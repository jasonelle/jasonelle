# file discovery, type classification, and corpus health checks
from __future__ import annotations
import fnmatch
import json
import os
import re
import shlex
import stat
import subprocess
import time
import unicodedata
from concurrent.futures import ThreadPoolExecutor
from enum import Enum
from functools import lru_cache
from pathlib import Path
from typing import Callable

from graphify.google_workspace import (
    GOOGLE_WORKSPACE_EXTENSIONS,
    convert_google_workspace_file,
    google_workspace_enabled,
)
from graphify.paths import GRAPHIFY_OUT, out_path


class FileType(str, Enum):
    CODE = "code"
    DOCUMENT = "document"
    PAPER = "paper"
    IMAGE = "image"
    VIDEO = "video"


_MANIFEST_PATH = str(out_path("manifest.json"))

#: Window in which a manifest row's own timestamp is too close to the file's
#: mtime for "mtime unchanged" to prove the content is unchanged. Coarse for
#: filesystems that round mtime to whole seconds; tight when real sub-second
#: precision is reported. Mirrors cache.py's `_MTIME_GRANULARITY_NS` (2s) —
#: the same racily-clean assumption at the hash-cache layer (#2466 / #2612);
#: keep the two in sync.
_MTIME_COARSE_S = 2.0
_MTIME_SUBSECOND_S = 0.05

CODE_EXTENSIONS = {'.py', '.ts', '.tsx', '.mts', '.cts', '.js', '.jsx', '.mjs', '.cjs', '.ejs', '.ets', '.go', '.rs', '.java', '.groovy', '.gradle', '.cpp', '.cc', '.cxx', '.c', '.h', '.hpp', '.cu', '.cuh', '.metal', '.rb', '.rake', '.swift', '.kt', '.kts', '.cs', '.scala', '.php', '.lua', '.luau', '.toc', '.zig', '.ps1', '.psm1', '.psd1', '.ex', '.exs', '.m', '.mm', '.ml', '.mli', '.jl', '.vue', '.svelte', '.astro', '.dart', '.v', '.sv', '.svh', '.sql', '.r', '.f', '.F', '.f90', '.F90', '.f95', '.F95', '.f03', '.F03', '.f08', '.F08', '.pas', '.pp', '.dpr', '.dpk', '.lpr', '.inc', '.dfm', '.lfm', '.lpk', '.sh', '.bash', '.json', '.tf', '.tfvars', '.hcl', '.dm', '.dme', '.dmi', '.dmm', '.dmf', '.sln', '.slnx', '.csproj', '.fsproj', '.vbproj', '.xaml', '.razor', '.cshtml', '.cls', '.trigger', '.lisp', '.cl', '.lsp', '.asd'}
DOC_EXTENSIONS = {'.md', '.mdx', '.qmd', '.skill', '.txt', '.rst', '.html', '.yaml', '.yml'}
PAPER_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg'}
OFFICE_EXTENSIONS = {'.docx', '.xlsx'}
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.webm', '.mkv', '.avi', '.m4v', '.mp3', '.wav', '.m4a', '.ogg'}

CORPUS_WARN_THRESHOLD = 50_000    # words - below this, warn "you may not need a graph"
CORPUS_UPPER_THRESHOLD = 500_000  # words - above this, warn about token cost
FILE_COUNT_UPPER = 500             # files - above this, warn about token cost

# Resource caps for parsing untrusted office/PDF files (F2). A corpus is
# attacker-controllable (graphify runs on cloned/shared folders), and .docx/.xlsx
# are zip+XML containers: a few-KB zip-bomb can decompress to gigabytes and
# OOM-kill the process at load_workbook/Document time. Screen the file before any
# parser touches it.
_OFFICE_MAX_RAW_BYTES = 50 * 1024 * 1024            # 50 MiB on-disk
_OFFICE_MAX_DECOMPRESSED_BYTES = 512 * 1024 * 1024  # 512 MiB total uncompressed
_OFFICE_MAX_COMPRESSION_RATIO = 200                 # uncompressed : compressed


def _file_within_size_cap(path: Path, cap: int = _OFFICE_MAX_RAW_BYTES) -> bool:
    """True if *path* exists and its on-disk size is within *cap*."""
    try:
        return path.stat().st_size <= cap
    except OSError:
        return False


def _zip_within_caps(path: Path) -> bool:
    """Reject a zip-based office file that is a likely zip/XML bomb.

    Two layers, because the zip central-directory sizes are attacker-controlled:
    1. A cheap pre-filter on the declared sizes (on-disk cap, summed-uncompressed
       cap, compression ratio) that rejects an honest bomb without decompressing.
    2. An authoritative pass that stream-decompresses every member with a hard
       byte ceiling, so a member that under-declares its size in the central
       directory cannot expand past the cap undetected. Decompression is chunked
       and bounded, so checking a bomb never materializes more than the ceiling.
    """
    import zipfile
    if not _file_within_size_cap(path):
        return False
    try:
        with zipfile.ZipFile(path) as zf:
            infos = zf.infolist()
            compressed = sum(i.compress_size for i in infos) or 1
            declared = sum(i.file_size for i in infos)
            if declared > _OFFICE_MAX_DECOMPRESSED_BYTES:
                return False
            if declared / compressed > _OFFICE_MAX_COMPRESSION_RATIO:
                return False
            total = 0
            for info in infos:
                with zf.open(info) as member:
                    while True:
                        chunk = member.read(1024 * 1024)
                        if not chunk:
                            break
                        total += len(chunk)
                        if total > _OFFICE_MAX_DECOMPRESSED_BYTES:
                            return False
    except (zipfile.BadZipFile, OSError, EOFError):
        return False
    return True

# Dedicated credential-store directories: everything beneath them is sensitive,
# with no carve-out — a .py inside ~/.ssh or ~/.aws is tooling for key material,
# not a source package, and keys there are routinely extensionless.
# Both sets are checked against path.parts[:-1] (parents only) so a root-level
# file named "credentials" or "secrets" is not falsely flagged by this stage.
_CREDENTIAL_STORE_DIRS = frozenset({
    ".ssh", ".gnupg", ".aws", ".gcloud",
})

# Bare-name directories that are as often legitimate source packages (Go
# internal/secrets, a credentials/ service module) as credential stores. Their
# contents are sensitive EXCEPT genuine programming-language source, mirroring
# the Stage 3 keyword carve-out (#1666) at the directory level (#1943).
_AMBIGUOUS_SENSITIVE_DIRS = frozenset({
    "secrets", ".secrets", "credentials",
})

# Files that may contain secrets - skip silently. These patterns are specific
# (extensions, exact credential-store names) and always apply.
_SENSITIVE_PATTERNS = [
    re.compile(r'(^|[\\/])\.(env|envrc)(\.|$)', re.IGNORECASE),
    re.compile(r'\.(pem|key|p12|pfx|cert|crt|der|p8)$', re.IGNORECASE),
    # SSH/GPG private keys. Left boundary + IGNORECASE so `grid_rsa` (alpha before
    # `id_rsa`) and `ID_RSA` are handled correctly, not matched as a substring.
    re.compile(r'(^|[^A-Za-z0-9])(id_rsa|id_dsa|id_ecdsa|id_ed25519)(\.pub)?$', re.IGNORECASE),
    re.compile(r'^secring(\.(gpg|pgp))?$', re.IGNORECASE),  # GPG private keyring
    # Auth/credential dotfiles that routinely hold tokens (#2106: .npmrc/.pypirc/
    # .git-credentials/.boto were silently indexed before).
    re.compile(r'(\.netrc|\.pgpass|\.htpasswd|\.npmrc|\.pypirc|\.git-credentials|\.boto)$', re.IGNORECASE),
    # NOTE: aws_credentials/gcloud_credentials/service_account moved to the
    # boundary-checked Stage 3 keyword path (#2106). The old unbounded
    # `service.account` substring (regex `.` wildcard) matched real source like
    # google/oauth2/service_account.py and prose like aws_credentials_rotation.md.
]

# Committed dotenv / envrc templates — placeholders only, not live secrets.
# Stage 2's `.env.` regex otherwise treats these like `.env.local` (#2184).
_ENV_TEMPLATE_SUFFIXES = (".example", ".sample", ".template", ".dist")


def _is_env_template(name: str) -> bool:
    """True for `.env.example` / `.envrc.sample` style committed templates (#2184)."""
    lower = name.lower()
    if not lower.endswith(_ENV_TEMPLATE_SUFFIXES):
        return False
    # Basename must still be an .env* / .envrc* file (not e.g. secrets.example).
    return bool(re.match(r"\.(env|envrc)\.", lower))

# Generic keyword patterns - these only count when the keyword is LOAD-BEARING
# in the filename (see _generic_keyword_hit), because a keyword buried mid-phrase
# in a long descriptive slug names a topic, not a credential store:
# "token-economics-of-recall.md" is a note ABOUT tokens; "api_token.txt" IS one.
# Uses lookarounds instead of \b so underscore-prefixed names like api_token.txt
# match. Both patterns use (?![a-zA-Z]) so that the trailing-underscore behavior
# is consistent: "secret_store.txt" IS flagged, "tokenizer.py" is NOT (because
# "i" after "token" is alpha and blocks the match).
# `token` is kept separate because its longer suffix "izer"/"ize" is the only
# common false-positive; other keywords have no such well-known derivatives.
_GENERIC_KEYWORD_PATTERNS = [
    re.compile(r'(?<![a-zA-Z0-9])(credential|secret|passwd|password|private_key)s?(?![a-zA-Z])', re.IGNORECASE),
    re.compile(r'(?<![a-zA-Z0-9])tokens?(?![a-zA-Z])', re.IGNORECASE),
    # service_account / service-account / serviceaccount (GCP key files). In the
    # keyword path so `service_account.py` (real source) is spared while
    # `service-account.json` (a downloaded key) and bare names are still caught
    # (#2106; was an unbounded Stage 2 substring). aws_credentials/gcloud_credentials
    # are already covered by the `credential` keyword above.
    re.compile(r'(?<![a-zA-Z0-9])service[._-]?account(?![a-zA-Z])', re.IGNORECASE),
]

# Prose/note formats: a heavily-linked wiki article whose topic slug ends in a
# keyword (privacy-tokens.md, token-economics.md) is a document ABOUT the topic,
# not a credential store, so it must not be silently dropped (#2106). A BARE
# keyword name (secrets.md, token.md, passwords.md) still reads as a dump and
# stays excluded — see _is_prose_note.
_PROSE_EXTS = frozenset({".md", ".markdown", ".rst", ".org", ".adoc", ".tex"})

# Data/serialization extensions that commonly ARE secret stores when their name
# hits a generic keyword (credentials.json, secrets.yaml, token.toml) or they sit
# in an ambiguous sensitive dir (secrets/db.json). These stay subject to the
# Stage 1 ambiguous-dir drop and the Stage 3 keyword drop even though some route
# through the CODE path for manifest parsing — only real programming-language
# source is exempt (#1666, #1943).
_SECRET_PRONE_DATA_EXTS = frozenset({
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".config",
    ".xml", ".properties", ".env", ".txt",
    # .tfvars is Terraform's canonical VALUES store (routinely holds real
    # secrets), not source — keep it out of the graph even though it sits in
    # CODE_EXTENSIONS. .tf/.hcl are genuine infra source and stay graphable.
    ".tfvars",
})

# Word separators for the load-bearing check (underscore intentionally included;
# multi-word keywords like private_key are handled by the end-of-stem check,
# which runs before word counting).
_WORD_SPLIT = re.compile(r'[-_\s.]+')  # '.' included so `token.economics.notes` counts as 3 words (#2106)


def _is_prose_note(path: Path) -> bool:
    """A prose/note file (.md/.rst/...) whose stem is a multi-word topic slug is
    exempt from the generic-keyword drop (#2106). A stem that IS exactly a bare
    keyword (secrets / token / passwords) is NOT exempt — that still reads as a
    credential dump."""
    if path.suffix.lower() not in _PROSE_EXTS:
        return False
    stem = Path(path.name).stem.lstrip('.') or Path(path.name).stem
    return not any(p.fullmatch(stem) for p in _GENERIC_KEYWORD_PATTERNS)


def _generic_keyword_hit(name: str) -> bool:
    """True if a generic secret keyword appears load-bearing in the filename.

    Secret-store files name their contents, and in English compounds the
    content noun is the head, which comes last: "github-personal-access-token",
    "api_token", "oauth_token". A keyword that is neither at the end of the
    stem nor in a short (<=2 word) name is a topic word in a descriptive slug
    ("token-economics-of-recall.md", "password-policy-discussion.md") and must
    not cause the file to be silently dropped from the graph (#436, #718).
    """
    # Stem = name minus only the FINAL extension (not up to the first dot), so a
    # multi-dot topic slug like `token.economics.notes.md` keeps all its words and
    # doesn't collapse to a bare `token` (#2106). Leading dots stripped so
    # dotfiles like `.token` keep their keyword.
    stem = Path(name).stem.lstrip('.') or Path(name).stem
    for pat in _GENERIC_KEYWORD_PATTERNS:
        hit = False
        for m in pat.finditer(stem):
            hit = True
            if m.end() == len(stem):  # keyword ends the stem -> names the contents
                return True
        if hit and len([w for w in _WORD_SPLIT.split(stem) if w]) <= 2:
            return True  # short name like token_config.yaml / secret_handler.txt
    return False

# Signals that a .md/.txt file is actually a converted academic paper
_PAPER_SIGNALS = [
    re.compile(r'\barxiv\b', re.IGNORECASE),
    re.compile(r'\bdoi\s*:', re.IGNORECASE),
    re.compile(r'\babstract\b', re.IGNORECASE),
    re.compile(r'\bproceedings\b', re.IGNORECASE),
    re.compile(r'\bjournal\b', re.IGNORECASE),
    re.compile(r'\bpreprint\b', re.IGNORECASE),
    re.compile(r'\\cite\{'),          # LaTeX citation
    re.compile(r'\[\d+\]'),           # Numbered citation [1], [23] (inline)
    re.compile(r'\[\n\d+\n\]'),       # Numbered citation spread across lines (markdown conversion)
    re.compile(r'eq\.\s*\d+|equation\s+\d+', re.IGNORECASE),
    re.compile(r'\d{4}\.\d{4,5}'),   # arXiv ID like 1706.03762
    re.compile(r'\bwe propose\b', re.IGNORECASE),   # common academic phrasing
    re.compile(r'\bliterature\b', re.IGNORECASE),   # "from the literature"
]
_PAPER_SIGNAL_THRESHOLD = 3  # need at least this many signals to call it a paper


def _is_graphable_source(path: Path) -> bool:
    """True for genuine programming-language source — the only category exempt
    from the ambiguous-dir (Stage 1, #1943) and generic-keyword (Stage 3, #1666)
    drops. Data/serialization formats are NOT exempt even though some route
    through the CODE path for manifest parsing: credentials.json / secrets.yaml
    are exactly the stores those stages must keep catching.
    """
    return classify_file(path) == FileType.CODE and path.suffix.lower() not in _SECRET_PRONE_DATA_EXTS


def _is_sensitive(path: Path) -> bool:
    """Return True if this file likely contains secrets and should be skipped."""
    # Stage 1: any PARENT directory is a known secrets dir (parts[:-1] excludes
    # the filename itself so a root-level file named "credentials" is not falsely
    # skipped — the name patterns in Stage 2 handle the filename). Dedicated
    # credential stores drop everything unconditionally; ambiguous bare-name dirs
    # (secrets/, credentials/) spare genuine source (#1943), which still falls
    # through so Stages 2-3 screen its filename like anywhere else.
    parents = path.parts[:-1]
    # Lowercase the segment comparison so `Secrets/`/`SECRETS/` (real on
    # case-insensitive macOS/Windows filesystems) are still caught (#2106).
    if any(part.lower() in _CREDENTIAL_STORE_DIRS for part in parents):
        return True
    if any(part.lower() in _AMBIGUOUS_SENSITIVE_DIRS for part in parents) and not _is_graphable_source(path):
        return True
    # Stage 2: filename pattern match. Template suffixes (.example/.sample/…)
    # on .env / .envrc are the usual "safe to commit" convention — keep them
    # in the graph without opening a broad Stage 2 allowlist (#2184 / #1921).
    name = path.name
    if any(p.search(name) for p in _SENSITIVE_PATTERNS) and not _is_env_template(name):
        return True
    # Stage 3: generic keywords, only when load-bearing in the name. Do NOT let a
    # bare name keyword silently drop a genuine programming-language source file:
    # a .rb/.py named device_token or passwords_controller is a module, not a secret
    # store (#1666). Data/config formats (.json, .yaml, .toml, ...) are deliberately
    # NOT exempt even though .json routes through the CODE path for manifest parsing,
    # because credentials.json / oauth_token.json / secrets.yaml are exactly the
    # secret stores this stage must catch. The specific Stage 2 patterns (.env, .pem,
    # id_rsa, ...) still apply to everything regardless of extension.
    if _generic_keyword_hit(name):
        # Genuine source AND multi-word prose notes are exempt; a bare-keyword
        # name (secrets.md, token.txt) still drops (#1666, #2106).
        return not (_is_graphable_source(path) or _is_prose_note(path))
    return False


def _looks_like_paper(path: Path) -> bool:
    """Heuristic: does this text file read like an academic paper?"""
    try:
        # Only scan first 3000 chars for speed
        text = path.read_text(encoding="utf-8", errors="ignore")[:3000]
        hits = sum(1 for pattern in _PAPER_SIGNALS if pattern.search(text))
        return hits >= _PAPER_SIGNAL_THRESHOLD
    except Exception:
        return False


_ASSET_DIR_MARKERS = {".imageset", ".xcassets", ".appiconset", ".colorset", ".launchimage"}


_SHEBANG_CODE_INTERPRETERS = {
    "python", "python3", "python2",
    "ruby", "perl", "node", "nodejs",
    "bash", "sh", "dash", "zsh", "fish", "ksh", "tcsh",
    "lua", "php", "julia", "Rscript",
}


def _split_env_s(value: str, rest: list[str]) -> list[str]:
    """Re-tokenize an `env -S`/`--split-string` packed command, prepending the
    operand to any trailing args. Returns the unpacked argv."""
    packed = " ".join([value, *rest]).strip()
    return shlex.split(packed)


def _env_command_args(args: list[str], *, allow_split: bool = True) -> list[str]:
    """Strip leading env(1) options and var assignments, return the trailing
    command argv. Covers macOS/BSD and GNU coreutils env documented spellings.

    POSIX/macOS short forms:
        env [-0iv] [-C workdir] [-P utilpath] [-S string]
            [-u name] [name=value ...] [utility [argument ...]]

    GNU coreutils long/compact forms additionally supported:
        --argv0=ARG / -a ARG / -aARG
        --unset=NAME / --unset NAME / -u NAME / -uNAME
        --chdir=DIR / --chdir DIR / -C DIR / -CDIR
        --split-string=STRING / --split-string STRING
        -S STRING / -SSTRING / -vS STRING / -vSSTRING
        --ignore-environment / --null / --debug / --list-signal-handling
        --default-signal[=SIG] / --ignore-signal[=SIG] / --block-signal[=SIG]

    `-S` / `--split-string` payloads are themselves env-style argument lists
    per the GNU shebang synopsis:
        #!/usr/bin/env -[v]S[option]... [name=value]... command [args]...
    so after splitting the payload we recursively re-parse it with
    `allow_split=False` (a nested -S inside a split payload is rejected to
    bound recursion).

    Unknown hyphen-prefixed args yield [] (we refuse to guess whether
    their next token is an interpreter or an operand).
    """
    i = 0
    while i < len(args):
        arg = args[i]

        if arg == "--":
            return args[i + 1:]

        # Split-string forms: tokenize the packed payload, then re-parse it
        # as env args (so leading assignments/flags inside the payload are
        # skipped before the interpreter is identified).
        if allow_split:
            if arg == "-S":
                if i + 1 >= len(args):
                    return []
                return _env_command_args(
                    _split_env_s(" ".join(args[i + 1:]), []),
                    allow_split=False,
                )
            if arg.startswith("-S") and len(arg) > 2:
                return _env_command_args(
                    _split_env_s(arg[2:], args[i + 1:]),
                    allow_split=False,
                )
            if arg == "-vS":
                if i + 1 >= len(args):
                    return []
                return _env_command_args(
                    _split_env_s(" ".join(args[i + 1:]), []),
                    allow_split=False,
                )
            if arg.startswith("-vS") and len(arg) > 3:
                return _env_command_args(
                    _split_env_s(arg[3:], args[i + 1:]),
                    allow_split=False,
                )
            if arg.startswith("--split-string="):
                return _env_command_args(
                    _split_env_s(arg.split("=", 1)[1], args[i + 1:]),
                    allow_split=False,
                )
            if arg == "--split-string":
                if i + 1 >= len(args):
                    return []
                return _env_command_args(
                    _split_env_s(args[i + 1], args[i + 2:]),
                    allow_split=False,
                )

        # Options with separate required operand
        if arg in {"-u", "-C", "-P", "-a", "--unset", "--chdir", "--argv0"}:
            if i + 2 > len(args):
                return []
            i += 2
            continue

        # Clumped short option + operand
        if (
            arg.startswith(("-u", "-C", "-P", "-a"))
            and len(arg) > 2
            and not arg.startswith("--")
        ):
            i += 1
            continue

        # Long option with `=` operand
        if arg.startswith(("--unset=", "--chdir=", "--argv0=")):
            i += 1
            continue

        # No-operand flags
        if arg in {"-", "-i", "-0", "-v", "--ignore-environment", "--null",
                   "--debug", "--list-signal-handling"}:
            i += 1
            continue

        # Signal-handling long flags (with or without =SIG operand — we treat
        # them as no-effect for interpreter-resolution purposes)
        if arg.startswith(("--default-signal", "--ignore-signal", "--block-signal")):
            i += 1
            continue

        # Unknown hyphen-prefixed: refuse to guess
        if arg.startswith("-"):
            return []

        # Inline NAME=value assignment
        if "=" in arg:
            i += 1
            continue

        # First non-option, non-assignment token starts the command argv
        return args[i:]

    return []


def _shebang_interpreter(path: Path) -> str | None:
    """Return the interpreter name from a shebang line.

    Handles forms that a naive parser misses:
      - `#!/usr/bin/env -S python3 -u`     (env -S split-args form, anywhere)
      - `#!/usr/bin/env -i bash`           (no-operand env flags)
      - `#!/usr/bin/env -u VAR python3`    (env options with operands)
      - `#!/usr/bin/env -C /tmp python3`   (env -C workdir)
      - `#!/usr/bin/env -P /bin python3`   (env -P utilpath)
      - `#!/usr/bin/env DEBUG=1 python3`   (inline var assignment)
      - `#!"/usr/local/bin/python with spaces"`  (shlex handles quotes)

    Returns the basename of the resolved interpreter, or None if there is
    no shebang / the file is unreadable / parsing fails.
    """
    try:
        with path.open("rb") as f:
            first = f.read(256)
        if not first.startswith(b"#!"):
            return None
        line = first.split(b"\n")[0].decode(errors="replace")[2:].strip()
        parts = shlex.split(line)
        if not parts:
            return None
        interp = Path(parts[0]).name
        if interp == "env":
            env_args = _env_command_args(parts[1:])
            if not env_args:
                return None
            interp = Path(env_args[0]).name
        return interp
    except (OSError, ValueError):
        return None


def _shebang_file_type(path: Path) -> FileType | None:
    """Peek at the first line of an extensionless file for a shebang."""
    interp = _shebang_interpreter(path)
    if interp in _SHEBANG_CODE_INTERPRETERS:
        return FileType.CODE
    return None


def classify_file(path: Path) -> FileType | None:
    # Package manifests (apm.yml, pyproject.toml, Cargo.toml, go.mod, pom.xml) are parsed
    # deterministically, so route them to the AST path (CODE) rather than the LLM
    # document path — otherwise apm.yml (a .yml "document") would be LLM-extracted
    # and a package would split into duplicate file-anchored nodes (#1377).
    from graphify.manifest_ingest import is_package_manifest_path
    if is_package_manifest_path(path):
        return FileType.CODE
    # Compound extensions must be checked before simple suffix lookup
    if path.name.lower().endswith(".blade.php"):
        return FileType.CODE
    ext = path.suffix.lower()
    if not ext:
        return _shebang_file_type(path)
    if ext in CODE_EXTENSIONS:
        return FileType.CODE
    if ext in PAPER_EXTENSIONS:
        # PDFs inside Xcode asset catalogs are vector icons, not papers
        if any(part.endswith(tuple(_ASSET_DIR_MARKERS)) for part in path.parts):
            return None
        return FileType.PAPER
    if ext in IMAGE_EXTENSIONS:
        return FileType.IMAGE
    if ext in DOC_EXTENSIONS:
        # Check if it's a converted paper
        if _looks_like_paper(path):
            return FileType.PAPER
        return FileType.DOCUMENT
    if ext in OFFICE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in GOOGLE_WORKSPACE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in VIDEO_EXTENSIONS:
        return FileType.VIDEO
    return None


def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF file using pypdf."""
    if not _file_within_size_cap(path):
        return ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""


def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text using python-docx."""
    if not _zip_within_caps(path):
        return ""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown text using openpyxl."""
    if not _zip_within_caps(path):
        return ""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            if len(rows) >= 1:
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                sections.extend([header, sep])
                for row in rows[1:]:
                    sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_extract_structure(path: Path) -> dict:
    """Extract structural nodes (sheets, named tables, column headers) from an .xlsx file.

    Returns a nodes/edges dict compatible with the graphify extract pipeline.
    Used in addition to xlsx_to_markdown so Claude sees both structure and content.
    """
    def _nid(*parts: str) -> str:
        return re.sub(r"[^a-z0-9_]", "_", "_".join(p.lower() for p in parts).strip("_"))

    try:
        import openpyxl
    except ImportError:
        return {"nodes": [], "edges": []}

    try:
        wb = openpyxl.load_workbook(str(path), read_only=False, data_only=True)
    except Exception:
        return {"nodes": [], "edges": []}

    # F-035: typo fix — was `_re.sub` (NameError, but unreachable because the
    # whole xlsx codepath is currently behind a feature flag / not yet wired
    # into the dispatcher). Before re-enabling this path, re-audit it for
    # zip/XML bombs (openpyxl is built on top of zipfile and lxml-style XML
    # parsing — a malicious .xlsx can blow up memory at load_workbook time).
    stem = re.sub(r"[^a-z0-9]", "_", path.stem.lower())
    str_path = str(path)
    file_nid = _nid(str_path)
    nodes: list[dict] = [{"id": file_nid, "label": path.name, "file_type": "document",
                           "source_file": str_path, "source_location": None}]
    edges: list[dict] = []
    seen: set[str] = {file_nid}

    def _add(nid: str, label: str) -> None:
        if nid not in seen:
            seen.add(nid)
            nodes.append({"id": nid, "label": label, "file_type": "document",
                           "source_file": str_path, "source_location": None})

    def _edge(src: str, tgt: str, relation: str) -> None:
        edges.append({"source": src, "target": tgt, "relation": relation,
                       "confidence": "EXTRACTED", "source_file": str_path,
                       "source_location": None, "weight": 1.0})

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_nid = _nid(stem, sheet_name)
        _add(sheet_nid, f"{sheet_name} (sheet)")
        _edge(file_nid, sheet_nid, "contains")

        # Named Excel Tables (ListObjects)
        if hasattr(ws, "tables"):
            for tbl in ws.tables.values():
                tbl_nid = _nid(stem, sheet_name, tbl.name)
                _add(tbl_nid, tbl.name)
                _edge(sheet_nid, tbl_nid, "contains")
                # Column headers from table header row
                ref = tbl.ref  # e.g. "A1:D10"
                if ref:
                    try:
                        from openpyxl.utils import range_boundaries
                        min_col, min_row, max_col, _ = range_boundaries(ref)
                        header_row = list(ws.iter_rows(min_row=min_row, max_row=min_row,
                                                       min_col=min_col, max_col=max_col,
                                                       values_only=True))
                        if header_row:
                            for col_name in header_row[0]:
                                if col_name:
                                    col_nid = _nid(stem, tbl.name, str(col_name))
                                    _add(col_nid, str(col_name))
                                    _edge(tbl_nid, col_nid, "contains")
                    except Exception:
                        pass
        else:
            # Fallback: first non-empty row as column headers
            for row in ws.iter_rows(max_row=1, values_only=True):
                for cell in row:
                    if cell:
                        col_nid = _nid(stem, sheet_name, str(cell))
                        _add(col_nid, str(cell))
                        _edge(sheet_nid, col_nid, "contains")
                break

    try:
        wb.close()
    except Exception:
        pass

    return {"nodes": nodes, "edges": edges}


def convert_office_file(path: Path, out_dir: Path, root: "Path | None" = None) -> Path | None:
    """Convert a .docx or .xlsx to a markdown sidecar in out_dir.

    Returns the path of the converted .md file, or None if conversion failed
    or the required library is not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None

    if not text.strip():
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    # Use a stable name derived from the original path to avoid collisions.
    # Hash the path RELATIVE to the scan root, not the absolute path: the
    # absolute form salts the name with the checkout location, so the same
    # tracked .xlsx in two clones/worktrees emits two differently-named,
    # byte-identical sidecars — unbounded duplicates when graphify-out/ is
    # committed, each ingested as a distinct source doc (#2059). The relative
    # path still disambiguates same-stem files in different directories.
    # Normalize to NFC before hashing: on macOS (HFS+/APFS) os.walk/rglob return
    # filenames in NFD, while Python string literals and directly-constructed
    # Path objects are NFC, so the same source file would otherwise hash to
    # different sidecar names across runs — making --update treat every Office
    # file as new and re-extract it (#1226).
    import hashlib
    import unicodedata
    if root is None:
        # Default layout: out_dir is <root>/<graphify-out>/converted.
        root = out_dir.parent.parent
    try:
        key = path.resolve().relative_to(Path(root).resolve()).as_posix()
    except (ValueError, OSError):
        # Not under the scan root (custom GRAPHIFY_OUT layouts, --include
        # sources, direct API callers): keep the previous absolute form rather
        # than guessing, so behavior is unchanged for those cases.
        key = str(path.resolve())
    normalized_path = unicodedata.normalize("NFC", key)
    name_hash = hashlib.sha256(normalized_path.encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    # Skip re-writing only when the sidecar is present AND at least as new as the
    # source. detect_incremental tracks the SIDECAR (not the Office source), so a
    # sidecar that is never rewritten after the source changes leaves the doc
    # reported "unchanged" forever and freezes the graph (#1649). Re-converting
    # when the source is newer bumps the sidecar's mtime/content, which the
    # incremental hash check then correctly picks up. An unchanged source keeps
    # its (newer-or-equal) sidecar untouched so it never churns (#1226).
    try:
        if out_path.exists() and os.stat(_os_path(out_path)).st_mtime >= os.stat(_os_path(path)).st_mtime:
            return out_path
    except OSError:
        if out_path.exists():
            return out_path
    out_path.write_text(
        f"<!-- converted from {path.name} -->\n\n{text}",
        encoding="utf-8",
    )
    return out_path


def count_words(path: Path) -> int:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return len(extract_pdf_text(path).split())
        if ext == ".docx":
            return len(docx_to_markdown(path).split())
        if ext == ".xlsx":
            return len(xlsx_to_markdown(path).split())
        # Only regular files may be opened. A repository can contain named
        # pipes, sockets and device nodes, and `clone <github-url>` exists to
        # point the scan at trees the operator did not write. open() on a FIFO
        # with no writer BLOCKS FOREVER — it never raises, so the except below
        # cannot help, and `graphify update` hangs with no output. os.stat
        # follows symlinks on purpose: a link pointing at a FIFO blocks exactly
        # like the FIFO itself.
        if not stat.S_ISREG(os.stat(_os_path(path)).st_mode):
            return 0
        with open(_os_path(path), encoding="utf-8", errors="ignore") as f:
            return len(f.read().split())
    except Exception:
        return 0


def _is_regular_file(path: Path) -> bool:
    """True only for regular files (symlinks followed).

    Named pipes, sockets and device nodes must never reach a reader:
    ``open()`` on a FIFO with no writer blocks forever and never raises, so a
    single ``pipe.py`` in a scanned repository hangs the run with no output.
    A symlink is resolved deliberately — a link pointing at a FIFO blocks
    exactly like the FIFO itself. A path that cannot be stat'ed is treated as
    not readable rather than raising.
    """
    try:
        return stat.S_ISREG(os.stat(path).st_mode)
    except OSError:
        return False


# Directory names to always skip - venvs, caches, build artifacts, deps
_SKIP_DIRS = {
    "venv", ".venv",  # "env"/".env"/"*_env" are gated on venv markers below (#2058)
    "node_modules", "__pycache__", ".git",
    "dist", "build", "target", "out",
    "site-packages", "lib64",
    ".pytest_cache", ".mypy_cache", ".ruff_cache",
    ".tox", ".nox", ".eggs", "*.egg-info",  # nox is tox's successor, same .nox/ venv shape (#1804)
    "graphify-out",  # never treat the default output as source input (#524)
    # Coverage/test-artefact dirs — generated, never architecturally meaningful
    "lcov-report",                          # Vitest/Istanbul/nyc HTML reports (#870);
                                            # bare "coverage" is gated on report
                                            # artefacts below (#2339)
    "visual-tests", "visual-test",          # Playwright/visual-regression bundles (#869)
    "__snapshots__",                        # Jest/Vitest snapshot dir (unambiguous)
    "storybook-static",                     # Storybook production build output
    "dist-protected",                       # Protected dist variants (same noise as dist)
    # Framework cache/build dirs — generated, never architecturally meaningful (#873)
    ".next", ".nuxt", ".turbo", ".angular",
    ".idea", ".cache", ".parcel-cache", ".svelte-kit", ".terraform", ".serverless",
    ".graphify",  # graphify's own extraction cache — never index self-generated data
    ".obsidian", ".smart-env",  # Obsidian vault metadata and plugin caches (#2493)
    ".worktrees",  # git worktree convention (#947) — sibling checkouts, always redundant
}

# Large generated files that are never useful to extract
_SKIP_FILES = {
    "package-lock.json", "yarn.lock", "pnpm-lock.yaml",
    "Cargo.lock", "poetry.lock", "Gemfile.lock",
    "composer.lock", "go.sum", "go.work.sum",
    # Removed allowlist config (#2112) — no longer consumed, so keep a leftover
    # file out of the unclassified list instead of surfacing it as scan input.
    ".graphifyinclude",
}

# A bare "snapshots" dir is a Jest/Vitest artifact only when it actually holds
# snapshot files or lives directly under a JS test root. Elsewhere it is often a
# real code namespace (e.g. Rails app/services/snapshots/), so pruning it by name
# silently dropped legitimate source from the graph (#1666). "__snapshots__" stays
# unconditionally pruned above; only the ambiguous bare name is gated here.
_JS_SNAPSHOT_TEST_ROOTS = frozenset({"__tests__", "__test__"})

# Files a coverage tool writes into its own output dir. Any one of them is proof
# the directory is generated: lcov (lcov.info), nyc/Istanbul (coverage-final.json,
# clover.xml, the lcov-report/ subtree), coverage.py (coverage.xml, .coverage),
# JaCoCo/Cobertura (jacoco.xml, cobertura-coverage.xml).
_COVERAGE_ARTIFACT_FILES = frozenset({
    "lcov.info", "coverage-final.json", "coverage-summary.json",
    "clover.xml", "coverage.xml", "cobertura-coverage.xml", "jacoco.xml",
    ".coverage", "index.html",
})
_COVERAGE_ARTIFACT_DIRS = frozenset({"lcov-report", "html-report"})


def _has_coverage_artifacts(d: "Path") -> bool:
    """True only when *d* holds files a coverage tool actually generated.

    ``coverage`` is a legitimate package name (a Python package, a Go/Rust module,
    a domain namespace), so pruning it by name alone silently drops real source —
    an entire 5-module package in #2339, with its dependents left in the graph so
    queries still returned plausible neighbours. Prune it only on real evidence,
    mirroring the ``snapshots``/``env`` gating (#1666/#2058): a coverage report
    file, or an Istanbul/lcov HTML report subtree.
    """
    try:
        for name in _COVERAGE_ARTIFACT_FILES:
            if (d / name).is_file():
                return True
        for name in _COVERAGE_ARTIFACT_DIRS:
            if (d / name).is_dir():
                return True
    except OSError:
        pass
    return False


def _has_venv_markers(d: "Path") -> bool:
    """True only when *d* has actual virtualenv/conda structure on disk.

    ``env``/``.env``/``*_env`` is a real source-directory convention (UVM/ASIC
    verification trees, and others), so pruning it by name alone silently drops
    legitimate source with no trace (#2058). Prune it only on real evidence: a
    ``pyvenv.cfg``, an ``activate`` script, a ``lib/python*`` tree, or conda's
    ``conda-meta/`` (``conda create -p ./env`` writes no pyvenv.cfg).
    """
    try:
        if (d / "pyvenv.cfg").is_file():
            return True
        if (d / "bin" / "activate").is_file() or (d / "Scripts" / "activate").is_file():
            return True
        if next(d.glob("lib/python*"), None) is not None:
            return True
        if (d / "conda-meta").is_dir():
            return True
    except OSError:
        pass
    return False


def _is_noise_dir(part: str, parent: "Path | None" = None) -> bool:
    """Return True if this directory name looks like a venv, cache, or dep dir."""
    if part in _SKIP_DIRS:
        return True
    if part in ("env", ".env") or part.endswith("_env"):
        # Ambiguous: a real venv OR a real source dir. Prune only on actual venv
        # evidence, mirroring the "snapshots" gating (#1666/#2058).
        if parent is None:
            return False  # cannot verify; keep a possibly-real code dir
        return _has_venv_markers(parent / part)
    if part == "coverage":
        # Ambiguous: a generated report dir OR a real package named coverage.
        # Prune only on actual coverage-artefact evidence (#2339).
        if parent is None:
            return False  # cannot verify; keep a possibly-real code dir
        return _has_coverage_artifacts(parent / part)
    if part == "snapshots":
        # Prune only when it looks like an actual JS/Vitest snapshot dir.
        if parent is None:
            return False  # cannot verify; keep a possibly-real code dir
        snap_dir = parent / part
        if parent.name in _JS_SNAPSHOT_TEST_ROOTS:
            return True
        try:
            if next(snap_dir.glob("*.snap"), None) is not None:
                return True
        except OSError:
            pass
        return False
    # Catch *_venv (unambiguous — "venv" is always a virtualenv signal). "*_env"
    # is gated on markers above (#2058), not pruned by name.
    if part.endswith("_venv"):
        return True
    if part.endswith(".egg-info"):
        return True
    # worktrees/ nested inside a dotted dir (e.g. .claude/worktrees/, .git/worktrees/)
    if part == "worktrees" and parent is not None and parent.name.startswith("."):
        return True
    return False


_VCS_MARKERS = (".git", ".hg", ".svn", "_darcs", ".fossil")


def _nfc(text: str) -> str:
    """Normalize text to NFC so ignore matching survives Unicode form drift.

    macOS (APFS/HFS+) returns filenames in NFD: "ç" comes back as "c" +
    U+0327 COMBINING CEDILLA. Editors write ignore files in NFC, where the
    same "ç" is the single codepoint U+00E7. The two render identically and
    compare unequal, so a pattern like `Orçamento/` silently fails to exclude
    the directory it names — the files are scanned and, for docs/PDFs, sent
    to an LLM despite an explicit rule against it.

    Both sides are normalized to NFC before any fnmatch call. NFC is the form
    Linux and Windows already use, so this is a no-op there and only repairs
    the macOS mismatch.
    """
    return unicodedata.normalize("NFC", text)


def _parse_gitignore_line(raw: str) -> str:
    """Parse one raw line from a .graphifyignore file per gitignore spec.

    - Strip newline chars
    - Strip inline comments (whitespace + # suffix), but only when # is
      preceded by whitespace — so path#with#hash.py is preserved
    - Unescape \\# to literal #
    - Remove trailing spaces unless escaped with backslash
    - Strip leading whitespace
    - Return empty string for blank lines and full-line comments
    """
    line = raw.rstrip("\n\r")
    line = line.lstrip()
    if not line or line.startswith("#"):
        return ""
    # Strip inline comments: require whitespace before # (gitignore extension)
    line = re.sub(r"\s+#+[^\\].*$", "", line)
    # Unescape \# → literal #
    line = line.replace("\\#", "#")
    # Remove unescaped trailing spaces (per gitignore spec)
    line = re.sub(r"(?<!\\) +$", "", line)
    return _nfc(line)


def _find_vcs_root(start: Path) -> Path | None:
    """Walk upward from start; return the first directory containing a VCS marker."""
    current = start.resolve()
    home = Path.home()
    while True:
        if any((current / m).exists() for m in _VCS_MARKERS):
            return current
        parent = current.parent
        if parent == current or current == home:
            return None
        current = parent


def _path_identity(path: Path) -> str:
    """Portable comparison key for an existing filesystem path."""
    return _nfc(os.path.normcase(os.path.abspath(os.fspath(path))))


def _git_tracked_path_keys(root: Path) -> tuple[set[str], set[str]]:
    """Return tracked-file keys and their ancestor-directory keys under *root*.

    Gitignore rules do not apply to paths already present in Git's index. Ask
    Git once per scan/predicate construction with NUL-delimited output so every
    valid filename is preserved. Missing Git, a non-Git VCS marker, command
    failure, and malformed output all fail closed to the historical ignore
    behavior rather than making discovery fail (#2759).
    """
    root = root.resolve()
    vcs_root = _find_vcs_root(root)
    if vcs_root is None or not (vcs_root / ".git").exists():
        return set(), set()
    try:
        proc = subprocess.run(
            ["git", "-C", str(vcs_root), "ls-files", "-z", "--cached"],
            stdin=subprocess.DEVNULL,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            check=False,
            timeout=30,
            env={**os.environ, "GIT_OPTIONAL_LOCKS": "0"},
        )
    except (OSError, subprocess.SubprocessError):
        return set(), set()
    if proc.returncode != 0:
        return set(), set()

    tracked_files: set[str] = set()
    tracked_dirs: set[str] = set()
    for raw in proc.stdout.split(b"\0"):
        if not raw:
            continue
        path = Path(os.path.abspath(vcs_root / os.fsdecode(raw)))
        try:
            path.relative_to(root)
        except ValueError:
            continue
        # Deleted index entries and submodule gitlinks are not discoverable
        # files. Symlinks to regular files remain eligible; the existing
        # in-root target guard still decides whether they may enter the corpus.
        if not _is_regular_file(path):
            continue
        tracked_files.add(_path_identity(path))
        parent = path.parent
        while parent != root:
            parent_key = _path_identity(parent)
            if parent_key in tracked_dirs:
                break  # its ancestors were added with the first file below it
            tracked_dirs.add(parent_key)
            parent = parent.parent
    return tracked_files, tracked_dirs


def _git_info_exclude(vcs_root: Path) -> Path | None:
    """Resolve ``$GIT_DIR/info/exclude`` for the repo rooted at ``vcs_root``.

    ``info/exclude`` is where git records local-only, uncommitted excludes — and
    where ``git worktree add`` writes nested worktree paths — so a repo can ignore
    a directory without any ``.gitignore`` entry. graphify only read
    ``.gitignore``/``.graphifyignore``, so it walked into those worktree copies and
    the graph exploded (#1810). Handles the linked-worktree/submodule case where
    ``.git`` is a file (``gitdir: <path>``) and the real excludes live in the
    shared common git dir. Returns None when there is no readable exclude file.
    """
    dot_git = vcs_root / ".git"
    git_dir: Path | None = None
    if dot_git.is_dir():
        git_dir = dot_git
    elif dot_git.is_file():
        try:
            content = dot_git.read_text(encoding="utf-8", errors="ignore").strip()
        except OSError:
            content = ""
        if content.startswith("gitdir:"):
            gd = Path(content[len("gitdir:"):].strip())
            if not gd.is_absolute():
                gd = (vcs_root / gd).resolve()
            git_dir = gd
            # A linked worktree's gitdir holds a `commondir` file pointing at the
            # shared git dir, where info/exclude actually lives.
            commondir = gd / "commondir"
            if commondir.exists():
                try:
                    cd_raw = commondir.read_text(encoding="utf-8", errors="ignore").strip()
                except OSError:
                    cd_raw = ""
                if cd_raw:
                    cd = Path(cd_raw)
                    git_dir = cd if cd.is_absolute() else (gd / cd).resolve()
    if git_dir is None:
        return None
    exclude = git_dir / "info" / "exclude"
    return exclude if exclude.is_file() else None


_warned_ignore_encodings: set[str] = set()


def _read_ignore_text(path: Path) -> str:
    """Read an ignore file, preferring UTF-8 but never silently dropping a rule.

    These files were read with ``errors="ignore"``, which turns a mis-encoded
    byte into *no* byte. An ignore file saved in the host's ANSI codepage — the
    historical Notepad default on Windows, and still what ``Set-Content`` writes
    without ``-Encoding`` — is not valid UTF-8, so ``Or\xe7amento/`` decoded to
    the pattern ``Oramento/``. That matches nothing, and nothing said so: the
    directory was scanned despite an explicit exclusion, which for a rule
    covering documents or PDFs means they reach the semantic pass anyway.

    So: UTF-8 (BOM-tolerant) first, since that is what the format should be and
    what every other reader here assumes. Only if that fails do we fall back to
    the host encoding, then to latin-1, which cannot fail and maps every byte to
    a codepoint — a rule spelled in some third encoding still comes out wrong,
    but it comes out *whole*, and the warning names the file so it is fixable.
    Decoding never raises, matching the previous contract.
    """
    raw = path.read_bytes()
    try:
        return raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        pass
    # A BOM'd UTF-16 file (common from PowerShell `Set-Content` / Notepad "Unicode")
    # is not valid UTF-8, and latin-1 would map its interleaved NUL bytes to a
    # wall of \x00, garbling every rule. Decode it as UTF-16 by its BOM first.
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        try:
            return raw.decode("utf-16")
        except UnicodeDecodeError:
            pass
    import locale
    import sys as _sys
    fallback = locale.getpreferredencoding(False) or "latin-1"
    for enc in (fallback, "latin-1"):
        try:
            text = raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
        key = str(path)
        if key not in _warned_ignore_encodings:
            _warned_ignore_encodings.add(key)
            print(
                f"[graphify] WARNING: {path} is not valid UTF-8; read it as "
                f"{enc} instead. Re-save it as UTF-8 — patterns with non-ASCII "
                "characters may not match as written.",
                file=_sys.stderr,
            )
        return text
    return raw.decode("utf-8", errors="ignore")


def _load_dir_own_ignore(d: Path, *, gitignore: bool = True) -> list[tuple[Path, str]]:
    """Read .gitignore/.graphifyignore directly inside *d* (not its ancestors).

    Merges .gitignore and .graphifyignore for this one directory (#1363):
    .gitignore is read first and .graphifyignore last, so .graphifyignore
    patterns (including `!` negations) win on conflict via last-match-wins;
    adding a .graphifyignore can only ever exclude MORE, never re-include a
    .gitignore-excluded file (#945 kept: a dir with only a .gitignore still
    gets sensible defaults).

    Shared by `_load_graphifyignore` (ancestor chain, loaded once before the
    scan) and the live os.walk loop in `detect()` (called per-directory as
    each descendant is visited), so nested ignore files *below* the scan
    root are honored too — previously only the scan root and its ancestors
    were read, so e.g. `vendor/sub/.gitignore` was silently ignored (#1206).
    """
    patterns: list[tuple[Path, str]] = []
    for fname in ((".gitignore", ".graphifyignore") if gitignore else (".graphifyignore",)):
        ignore_file = d / fname
        if ignore_file.exists():
            for raw in _read_ignore_text(ignore_file).splitlines():
                line = _parse_gitignore_line(raw)
                if line:
                    patterns.append((d, line))
    return patterns


def _load_graphifyignore(root: Path, *, gitignore: bool = True) -> list[tuple[Path, str]]:
    """Read .graphifyignore files and return (anchor_dir, pattern) pairs.

    Patterns are returned outer-first so that inner (closer) rules are
    appended last and win via last-match-wins semantics — matching gitignore
    behavior exactly.

    Walk ceiling: the nearest VCS root if inside a repo, otherwise the scan
    root itself (hermetic — no leakage across unrelated sibling projects).

    Covers the scan root and its ancestors only — directories *below* the
    scan root are picked up live during the os.walk in `detect()` instead,
    since they aren't known until the walk reaches them (#1206).
    """
    root = root.resolve()
    ceiling = _find_vcs_root(root) or root

    # Collect ancestor dirs from ceiling down to root (outer → inner)
    dirs: list[Path] = []
    current = root
    while True:
        dirs.append(current)
        if current == ceiling:
            break
        current = current.parent
    dirs.reverse()  # ceiling first, scan root last

    patterns: list[tuple[Path, str]] = []

    # $GIT_DIR/info/exclude is repo-root-scoped and, per git, ranks below every
    # per-directory .gitignore/.graphifyignore — so load it first (lowest priority
    # under last-match-wins) anchored at the VCS root, letting a nearer `!`
    # re-include still override it (#1810).
    info_exclude = _git_info_exclude(ceiling) if gitignore else None
    if info_exclude is not None:
        for raw in _read_ignore_text(info_exclude).splitlines():
            line = _parse_gitignore_line(raw)
            if line:
                patterns.append((ceiling, line))

    for d in dirs:
        patterns.extend(_load_dir_own_ignore(d, gitignore=gitignore))
    return patterns


def _match_anchored_ignore_pattern(path: str, pattern: str) -> bool:
    """Match an anchored gitignore pattern without letting ``*`` cross ``/``."""
    path_parts = tuple(path.split("/"))
    pattern_parts = tuple(pattern.split("/"))

    @lru_cache(maxsize=None)
    def _matches(path_idx: int, pattern_idx: int) -> bool:
        if pattern_idx == len(pattern_parts):
            return path_idx == len(path_parts)

        part = pattern_parts[pattern_idx]
        if part == "**":
            if pattern_idx == len(pattern_parts) - 1:
                return path_idx < len(path_parts)
            return _matches(path_idx, pattern_idx + 1) or (
                path_idx < len(path_parts)
                and _matches(path_idx + 1, pattern_idx)
            )

        return (
            path_idx < len(path_parts)
            and fnmatch.fnmatchcase(path_parts[path_idx], part)
            and _matches(path_idx + 1, pattern_idx + 1)
        )

    return _matches(0, 0)


def _is_ignored(
    path: Path,
    root: Path,
    patterns: list[tuple[Path, str]],
    *,
    _cache: dict[Path, bool] | None = None,
) -> bool:
    """Return True if the path should be ignored per .graphifyignore patterns.

    Uses gitignore last-match-wins semantics: all patterns are evaluated in
    order; the final matching pattern determines the result. Negation patterns
    (starting with !) un-ignore a previously ignored path.

    Enforces gitignore's parent-exclusion rule: a ! pattern cannot re-include
    a file whose ancestor directory is already excluded.

    _cache: optional dict shared across calls within the same scan. Ancestor
    directory results are memoised so files under the same subtree don't
    re-evaluate the same patterns repeatedly.
    """
    if not patterns:
        return False

    def _eval(target: Path) -> bool:
        """Apply last-match-wins to a single target path."""
        if _cache is not None and target in _cache:
            return _cache[target]
        def _matches(rel: str, p: str, path_relative: bool) -> bool:
            if path_relative:
                return _match_anchored_ignore_pattern(rel, p)
            parts = rel.split("/")
            if fnmatch.fnmatch(rel, p):
                return True
            if fnmatch.fnmatch(_nfc(target.name), p):
                return True
            for i, part in enumerate(parts):
                if fnmatch.fnmatch(part, p):
                    return True
                if fnmatch.fnmatch("/".join(parts[:i + 1]), p):
                    return True
            return False

        result = False
        for anchor, pattern in patterns:
            negated = pattern.startswith("!")
            raw = pattern[1:] if negated else pattern
            directory_only = raw.endswith("/")
            path_relative = "/" in raw.rstrip("/")
            p = raw.strip("/")
            if not p:
                continue

            # gitignore semantics: patterns from A/.gitignore apply ONLY to paths
            # under A. Matching non-anchored patterns against root-relative paths
            # let e.g. .hypothesis/.gitignore's bare "*" ignore the ENTIRE repo
            # (detect() returned 0 files). The anchor dir itself is exempt — an
            # ignore file governs its directory's contents, not the directory.
            matched = False
            try:
                rel_anchor = _nfc(str(target.relative_to(anchor)).replace(os.sep, "/"))
            except ValueError:
                continue  # target outside this pattern's anchor: cannot match
            if rel_anchor != ".":
                rel = rel_anchor
                if not path_relative:
                    try:
                        if len(root.parts) > len(anchor.parts):
                            rel = _nfc(str(target.relative_to(root)).replace(os.sep, "/"))
                    except ValueError:
                        pass
                matched = _matches(rel, p, path_relative=path_relative)
                if matched and directory_only and not target.is_dir():
                    matched = False

            if matched:
                result = not negated  # last match wins; ! flips to un-ignore
        if _cache is not None:
            _cache[target] = result
        return result

    # Gitignore parent-exclusion rule: a ! re-include cannot rescue a file
    # whose ancestor directory is already excluded. Walk ancestors top-down;
    # if any ancestor is excluded, the file is excluded regardless of later
    # ! patterns targeting the file or a sub-path.
    try:
        rel_parts = path.relative_to(root).parts
    except ValueError:
        return _eval(path)

    ancestor = root
    for part in rel_parts[:-1]:
        ancestor = ancestor / part
        if _eval(ancestor):
            return True
    return _eval(path)


def _is_scan_ignored(
    path: Path,
    root: Path,
    patterns: list[tuple[Path, str]],
    explicit_patterns: list[tuple[Path, str]],
    tracked_files: set[str],
    tracked_dirs: set[str],
    *,
    cache: dict[Path, bool],
    explicit_cache: dict[Path, bool],
) -> bool:
    """Apply ignore rules while preserving Git-tracked paths.

    ``patterns`` combines Git and graph-specific rules. ``explicit_patterns``
    contains only .graphifyignore/--exclude rules, which remain authoritative
    even for tracked files. A tracked file's ancestor directories are preserved
    from Git-only pruning so the walk can reach the file (#2759).
    """
    if not _is_ignored(path, root, patterns, _cache=cache):
        return False
    if _is_ignored(path, root, explicit_patterns, _cache=explicit_cache):
        return True
    identity = _path_identity(path)
    return identity not in tracked_files and identity not in tracked_dirs


def ignored_predicate(
    root: Path,
    *,
    extra_excludes: list[str] | None = None,
    gitignore: bool = True,
) -> Callable[[Path], bool]:
    """Build a per-path predicate answering "would detect() exclude this path?".

    Mirrors detect()'s ignore decisions for a single existing path WITHOUT
    re-walking the corpus, from the same machinery detect() uses: the ancestor
    .graphifyignore/.gitignore chain (_load_graphifyignore), CLI/persisted
    ``--exclude`` patterns appended last at the root anchor (#947), nested
    per-directory ignore files along the path's own lineage (#1206), the
    _is_noise_dir directory pruning, and _SKIP_FILES. The sensitive-file
    heuristic (_is_sensitive) is deliberately NOT included: callers use this
    predicate as positive evidence of a live ignore RULE (#2495), and a
    heuristic match is not user intent.

    Nested patterns are loaded lazily, once per directory, into one shared
    pattern list. That accumulation cannot cross-contaminate results — a
    pattern only ever matches paths under its anchor directory, so patterns
    from a sibling subtree are inert — which is the same invariant detect()'s
    live os.walk relies on, and it keeps the shared _is_ignored cache valid.
    """
    root = root.resolve()
    patterns = _load_graphifyignore(root, gitignore=gitignore)
    explicit_patterns = _load_graphifyignore(root, gitignore=False)
    # Only shell out to git when .gitignore actually contributes patterns beyond
    # the explicit (.graphifyignore/--exclude) set: with no .gitignore in play,
    # nothing is dropped by gitignore and the tracked-file exemption is moot, so a
    # non-.gitignore corpus pays no `git ls-files` cost.
    tracked_files, tracked_dirs = (
        _git_tracked_path_keys(root)
        if gitignore and len(patterns) > len(explicit_patterns)
        else (set(), set())
    )
    if extra_excludes:
        for pat in extra_excludes:
            line = _parse_gitignore_line(pat)
            if line:
                patterns.append((root, line))
                explicit_patterns.append((root, line))
    cache: dict[Path, bool] = {}
    explicit_cache: dict[Path, bool] = {}
    # root's own ignore file is the last entry of _load_graphifyignore's chain.
    loaded_dirs: set[Path] = {root}

    def _ignored(path: Path) -> bool:
        path = Path(os.path.abspath(path))
        try:
            rel_parts = path.relative_to(root).parts
        except ValueError:
            return False  # outside the scan root: detect() never considered it
        if path.name in _SKIP_FILES:
            return True
        # Noise-dir pruning: os.walk never descends these, so anything beneath
        # one is excluded from the corpus regardless of ignore patterns.
        parent = root
        for part in rel_parts[:-1]:
            if _is_noise_dir(part, parent):
                return True
            parent = parent / part
        # Load ignore files along this path's own lineage — detect()'s walk
        # would have loaded exactly these before reaching the file (#1206).
        ancestor = root
        for part in rel_parts[:-1]:
            ancestor = ancestor / part
            if ancestor not in loaded_dirs:
                loaded_dirs.add(ancestor)
                patterns.extend(_load_dir_own_ignore(ancestor, gitignore=gitignore))
                explicit_patterns.extend(
                    _load_dir_own_ignore(ancestor, gitignore=False)
                )
        return _is_scan_ignored(
            path,
            root,
            patterns,
            explicit_patterns,
            tracked_files,
            tracked_dirs,
            cache=cache,
            explicit_cache=explicit_cache,
        )

    return _ignored


def _auto_follow_symlinks(root: Path) -> bool:
    """Return whether ``root`` has any direct symlinked child.

    Kept for callers that import the private helper, but detection no longer
    enables symlink following automatically. Following symlinks is now an
    explicit opt-in, and out-of-root symlink targets are never indexed.
    """
    try:
        for p in root.iterdir():
            if p.is_symlink():
                return True
    except (OSError, PermissionError):
        pass
    return False


def _resolves_under_root(path: Path, root: Path) -> bool:
    """True when ``path`` resolves to a target inside ``root``."""
    try:
        path.resolve().relative_to(root.resolve())
    except (OSError, RuntimeError, ValueError):
        return False
    return True


def detect(root: Path, *, follow_symlinks: bool | None = None, google_workspace: bool | None = None, extra_excludes: list[str] | None = None, cache_root: Path | None = None, gitignore: bool = True) -> dict:
    root = root.resolve()
    configured_out_dir = root / GRAPHIFY_OUT
    configured_out_names = {configured_out_dir.name}
    try:
        configured_out_dir = configured_out_dir.resolve()
    except (OSError, RuntimeError):
        configured_out_dir = configured_out_dir.absolute()
    configured_out_names.add(configured_out_dir.name)
    # .graphifyinclude support was removed (#2112): its loader and matchers had
    # no consumers, so the file has been a silent no-op since dot directories
    # became indexed by default (#873). Surface that once per scan so a
    # leftover allowlist file is not a silent behavior change.
    if (root / ".graphifyinclude").is_file():
        import sys as _sys
        print(
            "[graphify] WARNING: .graphifyinclude is no longer supported "
            "(it has been non-functional since dot directories became indexed "
            "by default); to re-include ignored paths, use ! negation patterns "
            "in .graphifyignore.",
            file=_sys.stderr,
        )
    if follow_symlinks is None:
        follow_symlinks = False
    google_workspace = google_workspace_enabled() if google_workspace is None else google_workspace
    files: dict[FileType, list[str]] = {
        FileType.CODE: [],
        FileType.DOCUMENT: [],
        FileType.PAPER: [],
        FileType.IMAGE: [],
        FileType.VIDEO: [],
    }
    total_words = 0

    def _wc(path: Path) -> int:
        # Cache word counts against each file's stat signature so unchanged
        # PDFs/docx aren't re-parsed on every run just to size the corpus (#1656).
        # cache_root (when given, e.g. from `extract --out`) keeps this cache out
        # of the scanned corpus (#1747).
        from graphify import cache as _cache
        return _cache.cached_word_count(path, root, count_words, cache_root=cache_root)

    skipped_sensitive: list[str] = []
    unclassified: list[str] = []
    # Files/dirs dropped by a .gitignore/.graphifyignore rule. Recorded so an
    # over-broad ignore (or a legitimately-ignored subtree) is visible instead
    # of silently vanishing from the graph (#1922). Directory-level entries keep
    # this bounded — a pruned `data/` is one entry, not one per contained file.
    ignored: list[str] = []
    pruned_noise: list[str] = []
    ignore_patterns = _load_graphifyignore(root, gitignore=gitignore)
    explicit_ignore_patterns = _load_graphifyignore(root, gitignore=False)
    # See ignored_predicate: skip the `git ls-files` subprocess when .gitignore
    # contributes no patterns, so a non-.gitignore corpus pays nothing for it.
    tracked_files, tracked_dirs = (
        _git_tracked_path_keys(root)
        if gitignore and len(ignore_patterns) > len(explicit_ignore_patterns)
        else (set(), set())
    )
    ignore_cache: dict[Path, bool] = {}  # shared across all _is_ignored calls in this scan
    explicit_ignore_cache: dict[Path, bool] = {}
    # CLI --exclude patterns are anchored at the scan root and appended last
    # so they win over any .graphifyignore/.gitignore rules (#947).
    if extra_excludes:
        for pat in extra_excludes:
            line = _parse_gitignore_line(pat)
            if line:
                ignore_patterns.append((root, line))
                explicit_ignore_patterns.append((root, line))

    def _ignored_for_scan(path: Path) -> bool:
        return _is_scan_ignored(
            path,
            root,
            ignore_patterns,
            explicit_ignore_patterns,
            tracked_files,
            tracked_dirs,
            cache=ignore_cache,
            explicit_cache=explicit_ignore_cache,
        )

    # Always include graphify-out/memory/ - query results filed back into the graph
    memory_dir = root / GRAPHIFY_OUT / "memory"
    scan_paths = [root]
    if memory_dir.exists():
        scan_paths.append(memory_dir)

    seen: set[Path] = set()
    all_files: list[Path] = []

    # os.walk swallows os.scandir errors by default (no onerror -> the failing
    # directory subtree is silently skipped). That turns a transient
    # PermissionError, or a directory created/deleted mid-walk (e.g. concurrent
    # writes racing the scan), into a partial file list and, downstream, a
    # silently partial graph.json. Record and surface every skipped directory
    # so an incomplete enumeration is visible rather than silent.
    walk_errors: list[str] = []

    def _on_walk_error(err: OSError) -> None:
        import sys as _sys
        target = getattr(err, "filename", None) or "<unknown>"
        walk_errors.append(f"{target}: {err}")
        print(
            f"[graphify] WARNING: could not scan {target} ({err}); "
            f"its files are missing from this run's enumeration.",
            file=_sys.stderr,
        )

    for scan_root in scan_paths:
        in_memory_tree = memory_dir.exists() and str(scan_root).startswith(str(memory_dir))
        for dirpath, dirnames, filenames in os.walk(
            scan_root, followlinks=follow_symlinks, onerror=_on_walk_error
        ):
            dp = Path(dirpath)
            if follow_symlinks and os.path.islink(dirpath):
                real = os.path.realpath(dirpath)
                parent_real = os.path.realpath(os.path.dirname(dirpath))
                if parent_real == real or parent_real.startswith(real + os.sep):
                    dirnames.clear()
                    continue
            if not in_memory_tree:
                # dp == root was already loaded by _load_graphifyignore (root is
                # the last entry in its ancestor chain); every other directory
                # reached by the walk is a descendant below the scan root, whose
                # own .gitignore/.graphifyignore is unknown until we get here.
                # Load it now, before pruning dp's children, so a nested ignore
                # file governs its own subtree the same way git honors it (#1206).
                if dp != root:
                    ignore_patterns.extend(_load_dir_own_ignore(dp, gitignore=gitignore))
                    explicit_ignore_patterns.extend(
                        _load_dir_own_ignore(dp, gitignore=False)
                    )
                # Prune noise dirs in-place so os.walk never descends into them.
                # Dot dirs are allowed — users often want .github/, .claude/, etc.
                # Framework caches (.next, .nuxt, …) are caught by _is_noise_dir.
                # Negations need no special-casing here: _is_ignored already applies
                # last-match-wins (so `!dir/` un-ignores a directory and it won't be
                # pruned) and the gitignore parent-exclusion rule (a `!` cannot rescue
                # a file beneath an excluded dir), so descending an ignored directory to
                # look for a re-included file is never necessary. The previous blanket
                # `has_negation` disabled directory pruning for EVERY ignored dir whenever
                # any `!` rule existed — e.g. a single `!docs/**` made the walk descend
                # bin/, obj/, wwwroot/, generated/, … : a pathological slowdown on large
                # repos for no correctness gain.
                kept_dirs: list[str] = []
                for d in dirnames:
                    child = dp / d
                    is_configured_out = False
                    if d in configured_out_names:
                        try:
                            is_configured_out = child.resolve() == configured_out_dir
                        except (OSError, RuntimeError):
                            pass
                    if is_configured_out:
                        pruned_noise.append(str(child) + os.sep)
                        continue
                    if _is_noise_dir(d, dp):
                        # Record pruned-as-noise dirs so a wrongly-pruned real
                        # source dir is at least traceable in the output rather
                        # than vanishing silently (#2058).
                        pruned_noise.append(str(dp / d) + os.sep)
                        continue
                    if _ignored_for_scan(dp / d):
                        ignored.append(str(dp / d) + os.sep)
                        continue
                    kept_dirs.append(d)
                dirnames[:] = kept_dirs
                if follow_symlinks:
                    safe_dirs: list[str] = []
                    for d in dirnames:
                        child = dp / d
                        if child.is_symlink() and not _resolves_under_root(child, root):
                            skipped_sensitive.append(str(child) + " [symlink target outside scan root]")
                            continue
                        safe_dirs.append(d)
                    dirnames[:] = safe_dirs
            for fname in filenames:
                if fname in _SKIP_FILES:
                    continue
                p = dp / fname
                if p not in seen:
                    seen.add(p)
                    all_files.append(p)

    all_files.sort(key=lambda p: str(p))

    out_base = Path(cache_root).resolve() if cache_root is not None else root
    converted_dir = out_base / GRAPHIFY_OUT / "converted"

    for p in all_files:
        # For memory dir files, skip hidden/noise filtering
        in_memory = memory_dir.exists() and str(p).startswith(str(memory_dir))
        if not in_memory:
            # Skip files inside our own converted/ dir (avoid re-processing sidecars)
            if str(p).startswith(str(converted_dir)):
                continue
        if not in_memory and _ignored_for_scan(p):
            ignored.append(str(p))
            continue
        if not _resolves_under_root(p, root):
            skipped_sensitive.append(str(p) + " [symlink target outside scan root]")
            continue
        if not _is_regular_file(p):
            # A repository may contain named pipes, sockets and device nodes,
            # and `clone <github-url>` exists precisely to point the scan at
            # trees the operator did not write.
            #
            # This has to be caught HERE, at the one place where a path is
            # admitted to the corpus, rather than at each read: the readers
            # number in the hundreds across the extractors, and open() on a
            # FIFO with no writer BLOCKS FOREVER — it never raises, so their
            # try/except cannot help and the whole run hangs with no output.
            skipped_sensitive.append(str(p) + " [not a regular file]")
            continue
        if _is_sensitive(p):
            skipped_sensitive.append(str(p))
            continue
        ftype = classify_file(p)
        if not ftype:
            # Considered but unclassifiable: an extension not in any supported set,
            # or an extensionless, non-shebang file (Dockerfile, Gemfile, Makefile,
            # Rakefile, LICENSE, ...). Previously these left no trace at all — not
            # counted, not listed — so a user couldn't tell they were seen (#1692).
            unclassified.append(str(p))
            continue
        if ftype:
            if p.suffix.lower() in GOOGLE_WORKSPACE_EXTENSIONS:
                if not google_workspace:
                    skipped_sensitive.append(
                        str(p)
                        + " [Google Workspace shortcut skipped - pass --google-workspace "
                        "or set GRAPHIFY_GOOGLE_WORKSPACE=1]"
                    )
                    continue
                try:
                    md_path = convert_google_workspace_file(p, converted_dir, xlsx_to_markdown=xlsx_to_markdown, root=root)
                except Exception as exc:
                    skipped_sensitive.append(str(p) + f" [Google Workspace export failed: {exc}]")
                    continue
                if md_path:
                    if _ignored_for_scan(md_path):
                        continue
                    files[ftype].append(str(md_path))
                    total_words += _wc(md_path)
                else:
                    skipped_sensitive.append(str(p) + " [Google Workspace export produced no readable text]")
                continue
            # Office files: convert to markdown sidecar so subagents can read them
            if p.suffix.lower() in OFFICE_EXTENSIONS:
                md_path = convert_office_file(p, converted_dir, root=root)
                if md_path:
                    if _ignored_for_scan(md_path):
                        continue
                    files[ftype].append(str(md_path))
                    total_words += _wc(md_path)
                else:
                    # Conversion failed (library not installed) - skip with note
                    skipped_sensitive.append(str(p) + " [office conversion failed - pip install graphifyy[office]]")
                continue
            files[ftype].append(str(p))
            if ftype != FileType.VIDEO:
                total_words += _wc(p)

    for ftype in files:
        files[ftype].sort()

    total_files = sum(len(v) for v in files.values())
    needs_graph = total_words >= CORPUS_WARN_THRESHOLD

    # Determine warning - lower bound, upper bound, or sensitive files skipped
    warning: str | None = None
    if not needs_graph:
        warning = (
            f"Corpus is ~{total_words:,} words - fits in a single context window. "
            f"You may not need a graph."
        )
    elif total_words >= CORPUS_UPPER_THRESHOLD or total_files >= FILE_COUNT_UPPER:
        warning = (
            f"Large corpus: {total_files} files · ~{total_words:,} words. "
            f"Semantic extraction will be expensive (many Claude tokens). "
            f"Consider running on a subfolder."
        )

    return {
        "files": {k.value: v for k, v in files.items()},
        "total_files": total_files,
        "total_words": total_words,
        "needs_graph": needs_graph,
        "warning": warning,
        "skipped_sensitive": skipped_sensitive,
        "unclassified": sorted(unclassified),
        "walk_errors": walk_errors,
        "ignored": sorted(ignored),
        "pruned_noise_dirs": sorted(pruned_noise),
        "graphifyignore_patterns": len(ignore_patterns),
        "scan_root": str(root.resolve()),
    }


def _os_path(path: Path) -> str:
    r"""Return an OS path string safe for open()/stat() on Windows long paths.

    On win32, paths longer than the legacy MAX_PATH (260 chars) are rejected by
    the plain file APIs unless prefixed with the extended-length marker ``\\?\``
    (which also requires a fully-qualified path). Without it, _md5_file /
    save_manifest / count_words silently fail to hash deeply-nested files, so
    their manifest entry never stabilizes and detect_incremental re-flags them
    as changed on every run (#1655). cache._normalize_path strips this prefix
    for stable KEYS; this adds it for I/O. Non-win32 and already-prefixed paths
    pass through unchanged.
    """
    import sys
    if sys.platform != "win32":
        return str(path)
    s = str(path)
    if s.startswith("\\\\?\\"):
        return s
    try:
        s = os.path.abspath(s)  # \\?\ requires a fully-qualified path
    except Exception:
        return str(path)
    if s.startswith("\\\\"):
        # UNC share \\server\share -> \\?\UNC\server\share
        return "\\\\?\\UNC\\" + s[2:]
    return "\\\\?\\" + s


def _md5_file(path: Path) -> str:
    """MD5 of file contents streamed in 64KB chunks — for change detection only."""
    import hashlib as _hl
    h = _hl.md5(usedforsecurity=False)
    try:
        with open(_os_path(path), "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
    except OSError:
        return ""
    return h.hexdigest()


def _stat_and_hash(path_str: str) -> tuple[str, float, str] | None:
    """Stat + MD5 a single file; returns None on OSError (e.g. deleted mid-run)."""
    try:
        p = Path(path_str)
        return path_str, os.stat(_os_path(p)).st_mtime, _md5_file(p)
    except OSError:
        return None


def _nfc(s: str) -> str:
    """NFC-normalize a path string used as a manifest key.

    On macOS, ``os.walk`` / ``getcwd`` yield NFD paths while path literals
    and many skill-substituted roots are NFC. Raw string compare then treats
    every file as both deleted and new, forcing a full re-extract (#2221).
    Same boundary as the Office sidecar hash fix (#1226).
    """
    import unicodedata
    return unicodedata.normalize("NFC", s)


def _to_relative_for_storage(key: str, root: Path) -> str:
    """Return ``key`` as a forward-slash relative path from ``root``.

    Keys outside ``root`` (out-of-tree symlinked sources, external --include
    paths) and already-relative keys pass through unchanged — mirrors the
    fallback in :func:`graphify.watch._relativize_source_files` so the
    on-disk artifact survives the round-trip even when some paths cannot be
    portably encoded.

    Only ``root`` is resolved — the key itself is relativized symbolically
    so an in-root symlink (e.g. ``alias.py -> sub/target.py``) is stored
    under its own name. Resolving the key would point the stored entry at
    the symlink target, and the original key would then miss on reload and
    re-extract on every incremental run.

    Both sides of ``relpath`` are NFC'd first: stamped keys may already be
    NFC while ``Path(root).resolve()`` is NFD on macOS, and a mixed-form
    compare would mark an in-root file as ``../…`` and keep it absolute
    (#2221 / #777).
    """
    p = Path(key)
    if not p.is_absolute():
        return key
    try:
        base = _nfc(str(Path(root).resolve()))
        rel = os.path.relpath(_nfc(str(p)), base)
    except (ValueError, OSError):
        return key  # outside root (e.g. Windows cross-drive)
    # ``os.path.relpath`` happily produces ``../foo`` for paths outside
    # root; mirror the prior ``relative_to``-raises-ValueError semantics by
    # keeping out-of-root entries in their absolute form.
    if rel == ".." or rel.startswith(".." + os.sep) or rel.startswith("../"):
        return key
    return rel.replace(os.sep, "/")


def _to_absolute_from_storage(key: str, root: Path) -> str:
    """Inverse of :func:`_to_relative_for_storage`.

    Re-anchor a stored key against ``root``. Already-absolute keys
    (legacy manifests, out-of-root entries) pass through unchanged so
    that newly-loaded manifests from before this change remain readable.
    Uses ``Path(root).resolve()`` so the produced absolute path matches
    what :func:`detect` returns (which also resolves the scan root).
    NFC both sides so a relative key and an NFD-resolved root still join
    to the same string form the rest of the manifest path uses (#2221).
    """
    p = Path(key)
    if p.is_absolute():
        return str(p)
    # NFC the joined result so an NFD-resolved root + relative key lands on
    # the same form load_manifest / detect_incremental compare against.
    return _nfc(str(Path(root).resolve() / p))


def load_manifest(
    manifest_path: str = _MANIFEST_PATH,
    *,
    root: Path | None = None,
) -> dict:
    """Load the manifest from a previous run. Returns {} on any error.

    When ``root`` is provided, stored relative keys are re-anchored against
    it so callers see absolute paths regardless of on-disk format. Legacy
    manifests with absolute keys pass through unchanged, so a graphify-out/
    written by an older version (or by a caller that didn't supply ``root``
    to :func:`save_manifest`) remains readable.

    Keys are NFC-normalized on load so a manifest written under one Unicode
    form still matches a scan that yields the other (#2221).
    """
    try:
        raw = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
    except Exception:
        return {}
    if not isinstance(raw, dict):
        return raw
    if root is None:
        return {_nfc(k): v for k, v in raw.items()}
    return {_nfc(_to_absolute_from_storage(k, root)): v for k, v in raw.items()}


def save_manifest(
    files: dict[str, list[str]],
    manifest_path: str = _MANIFEST_PATH,
    *,
    kind: str = "both",
    root: Path | None = None,
    scan_corpus: set[str] | list[str] | None = None,
    clear_semantic: set[str] | list[str] | None = None,
    clear_ast: set[str] | list[str] | None = None,
) -> None:
    """Save current file mtimes + content hashes for change detection.

    kind="ast"      — written by `graphify update` (AST-only rebuild). Stamps
                      ast_hash; preserves an existing semantic_hash only when
                      the file content is unchanged (mtime + hash match).
    kind="semantic" — written by `graphify extract` after semantic extraction.
                      Stamps semantic_hash; preserves existing ast_hash.
    kind="both"     — full pipeline: stamps both hashes (default).

    When ``root`` is provided, keys are relativized against it before write
    (forward-slash, posix-style) so the on-disk manifest is portable across
    machines and checkout locations (#777). Out-of-root entries are written
    as absolute so they continue to round-trip on the saving machine.
    When ``root`` is None the legacy absolute-keyed format is preserved.

    ``scan_corpus`` (#1908): full-scan callers pass the COMPLETE detect
    corpus (absolute paths) so seeded rows for in-root files that are still
    alive on disk but no longer part of the scan (newly excluded via
    .graphifyignore/.gitignore/--exclude) are dropped instead of surviving
    forever and masquerading as deletions in detect_incremental. It must be
    the RAW detect output, not a stamp-filtered subset — pruning to a
    filtered set would erase rows the filter merely omitted (failed chunks,
    --code-only doc rows). Out-of-root entries are never pruned. Callers
    saving a SUBSET of files (changed_paths hooks, skill runbooks, #917)
    must leave this None so their untouched rows are preserved.

    ``clear_semantic`` (#1948): files that were dispatched this run but
    produced no stamped output (e.g. the LLM omitted their chunk on a
    --force re-run) are absent from ``files``, so the seed loop below would
    otherwise copy their prior semantic_hash verbatim — masking the omission
    and making detect_incremental(kind="semantic") report them unchanged.
    Pass the set of such files (any path form ``scan_corpus`` accepts) to
    force their seeded semantic_hash to "" instead of inheriting it.

    ``clear_ast`` (#2543): same idea for AST failures (missing optional extra,
    zero-node anomalous extract). Blanks BOTH ``ast_hash`` and
    ``semantic_hash`` on the seeded row so either detect_incremental kind
    re-queues the file after the failure is fixed, without deleting
    graphify-out/.
    """
    existing = load_manifest(manifest_path, root=root)

    # Index both raw and NFC forms so scan/clear membership survives the
    # same NFC/NFD mismatch that breaks manifest lookups (#2221).
    def _path_index(paths: set[str] | list[str] | None) -> set[str] | None:
        if paths is None:
            return None
        indexed: set[str] = set()
        for p in paths:
            indexed.add(p)
            indexed.add(_nfc(p))
        return indexed

    scan_set = _path_index(scan_corpus)
    clear_set = _path_index(clear_semantic)
    clear_ast_set = _path_index(clear_ast)
    try:
        root_res: Path | None = Path(root).resolve() if root is not None else None
    except (OSError, RuntimeError):
        root_res = Path(root) if root is not None else None

    def _in_scan(path_str: str) -> bool:
        if path_str in scan_set or _nfc(path_str) in scan_set:
            return True
        try:
            resolved = str(Path(path_str).resolve())
            return resolved in scan_set or _nfc(resolved) in scan_set
        except (OSError, RuntimeError):
            return False

    def _in_clear(path_str: str) -> bool:
        if clear_set is None:
            return False
        if path_str in clear_set or _nfc(path_str) in clear_set:
            return True
        try:
            resolved = str(Path(path_str).resolve())
            return resolved in clear_set or _nfc(resolved) in clear_set
        except (OSError, RuntimeError):
            return False

    def _in_clear_ast(path_str: str) -> bool:
        if clear_ast_set is None:
            return False
        if path_str in clear_ast_set or _nfc(path_str) in clear_ast_set:
            return True
        try:
            resolved = str(Path(path_str).resolve())
            return resolved in clear_ast_set or _nfc(resolved) in clear_ast_set
        except (OSError, RuntimeError):
            return False

    def _in_root(path_str: str) -> bool:
        # Without a root we cannot tell in-root from out-of-root; fail open
        # (keep the row) so out-of-root corpora are never pruned by accident.
        if root_res is None:
            return False
        p = Path(path_str)
        try:
            p.relative_to(root_res)
            return True
        except ValueError:
            pass
        try:
            p.resolve().relative_to(root_res)
            return True
        except (ValueError, OSError, RuntimeError):
            return False

    def _normalise_entry(entry):
        if isinstance(entry, (int, float)):
            return {"mtime": entry, "ast_hash": "", "semantic_hash": ""}
        if isinstance(entry, dict) and "hash" in entry and "ast_hash" not in entry:
            return {"mtime": entry.get("mtime", 0), "ast_hash": entry["hash"], "semantic_hash": ""}
        if isinstance(entry, dict):
            return entry
        return None

    # Seed from the existing manifest so incremental callers passing a subset
    # of files don't silently erase entries for untouched files (#917).
    # Prune entries whose file no longer exists on disk — those are genuine
    # deletions that detect_incremental() should treat as gone. When the
    # caller supplied the full scan corpus, additionally prune in-root rows
    # the scan no longer covers: those files were excluded, not deleted, and
    # keeping the row makes them look deleted on every future run (#1908).
    manifest: dict[str, dict] = {}
    for f, entry in existing.items():
        normalised = _normalise_entry(entry)
        if normalised is None:
            continue
        try:
            if not Path(f).exists():
                continue
        except OSError:
            continue
        if scan_set is not None and not _in_scan(f) and _in_root(f):
            continue  # excluded-but-alive: drop the stale row (#1908)
        if clear_ast_set is not None and _in_clear_ast(f):
            # AST failure this run (missing extra / zero nodes, #2543): blank
            # both hashes so either detect_incremental kind re-queues.
            normalised = {**normalised, "ast_hash": "", "semantic_hash": ""}
        elif clear_set is not None and _in_clear(f):
            # Dispatched-but-omitted this run: don't inherit the stale
            # semantic_hash, or detect_incremental would call it unchanged (#1948).
            normalised = {**normalised, "semantic_hash": ""}
        manifest[f] = normalised

    all_files = [f for file_list in files.values() for f in file_list]
    with ThreadPoolExecutor() as pool:
        raw = pool.map(_stat_and_hash, all_files)
    hashed: dict[str, tuple[float, str]] = {
        r[0]: (r[1], r[2]) for r in raw if r is not None
    }

    for f in all_files:
        if f not in hashed:
            continue  # file deleted between detect() and manifest write
        mtime, h = hashed[f]
        key = _nfc(f)
        prev = _normalise_entry(existing.get(key, {})) or {}
        if kind in ("ast", "both"):
            ast_h = h
        else:
            ast_h = prev.get("ast_hash", "")
        if kind in ("semantic", "both"):
            sem_h = h
        else:
            # Preserve semantic_hash only when content is unchanged
            sem_h = prev.get("semantic_hash", "") if h == prev.get("ast_hash", "") else ""

        # Preserve previous seen timestamp if the entry's mtime and target hash(es)
        # are genuinely unchanged and no clear was requested for this file.
        prev_seen = prev.get("seen")
        is_unchanged = (
            isinstance(prev_seen, (int, float))
            and mtime == prev.get("mtime")
            and (ast_h == prev.get("ast_hash", "") if kind in ("ast", "both") else True)
            and (sem_h == prev.get("semantic_hash", "") if kind in ("semantic", "both") else True)
            and not _in_clear_ast(f)
            and not _in_clear(f)
        )
        entry: dict = {
            "mtime": mtime,
            "seen": prev_seen if is_unchanged else time.time(),
            "ast_hash": ast_h,
            "semantic_hash": sem_h,
        }
        manifest[key] = entry
    if root is not None:
        # Persist in portable form: forward-slash relative paths. Keys outside
        # ``root`` (out-of-tree symlinked corpora, --include sources) keep
        # their absolute form so the manifest round-trips on the saving
        # machine even when not every entry can be portably encoded.
        # NFC after relativize so on-disk keys match what load_manifest
        # re-anchors and compares against (#2221).
        manifest = {_nfc(_to_relative_for_storage(k, root)): v for k, v in manifest.items()}
    else:
        manifest = {_nfc(k): v for k, v in manifest.items()}

    # Avoid rewriting manifest.json when the serialized payload is identical (#2838).
    manifest_p = Path(manifest_path)
    if manifest_p.is_file():
        try:
            disk_raw = json.loads(manifest_p.read_text(encoding="utf-8"))
            if isinstance(disk_raw, dict) and disk_raw == manifest:
                return
        except Exception:
            pass

    from graphify.paths import write_json_atomic
    # Atomic write: a crash mid-write must not leave a truncated manifest that
    # detect_incremental then fails to parse.
    write_json_atomic(manifest_path, manifest, indent=2)


def _mtime_may_hide_a_rewrite(current_mtime: float, stored: dict) -> bool:
    """Was this manifest row written in the same tick as the file it describes?

    The incremental gate treats "mtime unchanged" as proof the content is
    unchanged. That is only true while the filesystem can distinguish the two
    writes: an edit keeping the file the same length and landing in the same
    timestamp tick moves neither size nor mtime, so the file silently skips
    re-extraction and the graph keeps serving the old content.

    ``seen`` records when the row was stamped. If the file's mtime falls inside
    the same tick, this row cannot prove currency and the caller pays for one
    MD5. Every other row — the whole settled corpus, and any manifest written
    by an earlier run — keeps the free stat-only fastpath.

    Rows predating ``seen`` are treated as safe: they necessarily come from an
    earlier process, where a later write would have had to move mtime.
    """
    seen = stored.get("seen")
    if not isinstance(seen, (int, float)):
        return False
    delta = float(seen) - float(current_mtime)
    if delta < 0:
        return False  # file is newer than the row; the mtime check already fired
    # Derive granularity from the timestamp: a whole-second mtime means the
    # filesystem cannot separate writes inside that second.
    coarse = float(current_mtime).is_integer()
    return delta < (_MTIME_COARSE_S if coarse else _MTIME_SUBSECOND_S)


def detect_incremental(
    root: Path,
    manifest_path: str = _MANIFEST_PATH,
    *,
    follow_symlinks: bool | None = None,
    google_workspace: bool | None = None,
    kind: str = "semantic",
    extra_excludes: list[str] | None = None,
    gitignore: bool = True,
) -> dict:
    """Like detect(), but returns only new or modified files since the last run.

    kind="semantic" (default for extract): a file is "changed" when its
        semantic_hash is missing or its content has changed since the last
        semantic extraction pass. Use this for `graphify extract` so that
        files touched by `graphify update` (AST-only) are re-extracted
        semantically.
    kind="ast": a file is "changed" when its ast_hash is missing or its
        content has changed. Use this for `graphify update`.

    Fast path: mtime unchanged + hash matches → unchanged (free, no disk IO
    beyond stat). Slow path: mtime bumped → compare MD5 against the relevant
    hash field before re-extracting.

    Backwards compatible with legacy manifests storing plain float mtime values
    or {mtime, hash} dicts (treated as ast_hash only; semantic_hash = miss).

    The ``follow_symlinks`` flag is forwarded to :func:`detect` so in-root
    symlinked sub-trees are scanned consistently between full and incremental
    runs. ``None`` (default) does not follow symlinked directories; callers must
    opt in explicitly, and resolved targets outside the scan root are skipped.
    """
    full = detect(
        root,
        follow_symlinks=follow_symlinks,
        google_workspace=google_workspace,
        extra_excludes=extra_excludes,
        gitignore=gitignore,
    )
    # Pass ``root`` so a manifest written with relative keys (post-#777) is
    # re-anchored to the absolute form the rest of this function compares
    # against. Legacy absolute-keyed manifests pass through unchanged.
    manifest = load_manifest(manifest_path, root=root)

    if not manifest:
        # No previous run - treat everything as new
        full["incremental"] = True
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_total"] = full["total_files"]
        full["deleted_files"] = []
        full["excluded_files"] = []
        return full

    new_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    unchanged_files: dict[str, list[str]] = {k: [] for k in full["files"]}

    for ftype, file_list in full["files"].items():
        for f in file_list:
            # Manifest keys are NFC; scan paths may arrive NFD (#2221).
            stored = manifest.get(_nfc(f))
            try:
                current_mtime = os.stat(_os_path(Path(f))).st_mtime
            except Exception:
                current_mtime = 0

            # Legacy manifest: plain float value stores only mtime.
            # Compare with `!=` so backwards mtime motion (git checkout of an
            # older commit, tarball restore, rsync --times) still triggers a
            # re-extract; the previous `>` silently kept the stale cache and
            # the graph drifted from disk (#1859). No stored hash means we
            # cannot verify content — any mtime delta forces a re-extract,
            # and the next save promotes the entry into the dict schema.
            if isinstance(stored, (int, float)):
                changed = current_mtime != stored
            elif isinstance(stored, dict):
                # Normalise legacy {mtime, hash} to new schema
                if "hash" in stored and "ast_hash" not in stored:
                    stored = {"mtime": stored.get("mtime", 0), "ast_hash": stored["hash"], "semantic_hash": ""}
                hash_key = "semantic_hash" if kind == "semantic" else "ast_hash"
                stored_hash = stored.get(hash_key, "")
                # Missing semantic_hash means update ran but extract hasn't — always re-extract
                if not stored_hash:
                    changed = True
                else:
                    stored_mtime = stored.get("mtime")
                    # Schema-drift guard (#1163): tolerate a nested {mtime: ...}
                    # dict or any non-numeric value without crashing.
                    if isinstance(stored_mtime, dict):
                        stored_mtime = stored_mtime.get("mtime")
                    if not isinstance(stored_mtime, (int, float)):
                        stored_mtime = None
                    if stored_mtime is None or current_mtime != stored_mtime:
                        # mtime bumped — verify with content hash before re-extracting
                        changed = _md5_file(Path(f)) != stored_hash
                    elif _mtime_may_hide_a_rewrite(current_mtime, stored):
                        # mtime is unchanged, but it was recorded in the same
                        # filesystem tick the file was written in — a later
                        # same-length edit lands in that tick without moving
                        # mtime, and the file silently skips re-extraction
                        # while the graph keeps serving the old content.
                        # Only this narrow window pays for a content hash.
                        changed = _md5_file(Path(f)) != stored_hash
                    else:
                        changed = False
            else:
                changed = True  # unknown format, re-extract to be safe

            if changed:
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    # Manifest rows that left the corpus, split by disk existence (#1908):
    # a row whose file is gone from DISK is a genuine deletion (its cached
    # nodes are ghosts); a row whose file still exists but is out of the
    # current scan was EXCLUDED (ignore rules / --exclude changed) and must
    # not be reported as deleted. Mirrors the watch-side excluded-vs-deleted
    # distinction (#1795).
    current_files = {_nfc(f) for flist in full["files"].values() for f in flist}
    deleted_files: list[str] = []
    excluded_files: list[str] = []
    for f in manifest:
        if _nfc(f) in current_files:
            continue
        try:
            alive = Path(f).exists()
        except OSError:
            alive = False
        (excluded_files if alive else deleted_files).append(f)

    new_total = sum(len(v) for v in new_files.values())
    full["incremental"] = True
    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_total"] = new_total
    full["deleted_files"] = deleted_files
    full["excluded_files"] = excluded_files
    return full
